import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

doc = Document()

# Page Margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("Nirvayu — Air Quality & Compliance Platform")
title_run.font.name = "Arial"
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(16, 149, 193) # #1095C1

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("Complete Feature Architecture, Technical Concepts & API Specifications Document")
sub_run.font.name = "Arial"
sub_run.font.size = Pt(13)
sub_run.font.italic = True
sub_run.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph() # Spacer

features = [
    {
        "num": "1",
        "title": "SafeLifePlanner (Personalized Citizen Air Safety Advisor)",
        "components": "client/src/components/SafeLifePlanner.tsx",
        "concept": "Multi-step React decision-tree state machine (landing -> onboarding -> result). Evaluates user session inputs (Age Group, Environment, Outdoor Exposure Hours, Today's Activity, Occupation) combined with live ward telemetry (AQI, PM2.5, PM10, NO2, dominant_source) to compute a personalized Environmental Exposure Concern level (LOW, MODERATE, HIGH, VERY HIGH), dynamic impact drivers, activity matrix, 3-5 prioritized recommendations, safer route exposure optimizer, and home window ventilation radar without persistent storage.",
        "api": "Reads live ward telemetry via GET /api/wards and 24h predicted AQI via ward.intelligence_data.predicted_aqi."
    },
    {
        "num": "2",
        "title": "Ward-Wise 7-Day Historic AQI Trend Graph",
        "components": "client/src/components/HistoricAqiChart.tsx, server/storage.ts",
        "concept": "Recharts AreaChart visualization rendering per-ward 7-day historic telemetry. Generates deterministic phase-shifted sinusoidal historical curves (buildHistoricAqi) anchored to live AQI, displaying KPI metrics for 7-Day Peak, 7-Day Low, and vs. Yesterday delta.",
        "api": "Served inside ward.intelligence_data.aqi_history via GET /api/wards."
    },
    {
        "num": "3",
        "title": "24-Hour Random Forest AQI Prediction Engine",
        "components": "server/rf_evaluator.ts, server/models/rf_trees.json, server/aqi_predictor.py, server/scripts/train_model.py",
        "concept": "Supervised Machine Learning Random Forest Regressor trained on 19 Delhi historical pollution CSV datasets (35 decision trees). Inlined statically into a zero-dependency TypeScript evaluator (rf_evaluator.ts) for <1ms execution on Vercel serverless functions, backed by a Python subprocess runner (aqi_predictor.py) and diurnal sine wave fallback.",
        "api": "POST /api/aqi/predict accepts { pm10, o3, no2, so2, co, timestamp } returning 24-hour forecast array { hourly: number[] }. Internal calls via predictFutureAqi() populate ward.intelligence_data.predicted_aqi."
    },
    {
        "num": "4",
        "title": "AI-Powered Pollution Image Verification & Classification",
        "components": "client/src/components/CaptureEvidence.tsx, server/routes.ts",
        "concept": "Multimodal Artificial Intelligence visual classification using Google Generative AI (Gemini Flash). Inspects base64 citizen photos for visible airborne pollutants (traffic, construction, stubble burning, other) and rejects non-pollution photos (clean sky, litter, water, selfies). Includes keyword fallback classifier when API keys or quotas are depleted.",
        "api": "Google Generative AI SDK (@google/generative-ai) via process.env.GEMINI_API_KEY (gemini-2.5-flash, gemini-2.0-flash, gemini-1.5-flash) invoked via POST /api/reports."
    },
    {
        "num": "5",
        "title": "Cryptographic & Blockchain-Backed Report Verification Ledger",
        "components": "server/blockchain.ts, client/src/lib/blockchain.ts",
        "concept": "Cryptographic SHA-256 content hashing (crypto.createHash('sha256')) combining image byte hashes, ward ID, pollution type, description, and timestamp into a 32-byte hex hash (mediaHash). Simulated on-chain immutability ledger persisting entries to attached_assets/reports_ledger.json.",
        "api": "POST /api/reports (creates SHA-256 hash & registers report) and GET /api/reports/blockchain-ledger (retrieves verified entries)."
    },
    {
        "num": "6",
        "title": "Geospatial Geofenced Evidence Verification",
        "components": "client/src/components/CaptureEvidence.tsx, server/routes.ts",
        "concept": "Geospatial spatial analysis using Turf.js (@turf/turf). Calculates distance (turf.distance) between citizen GPS coordinates and ward centroids to enforce physical presence within a 3.0 km boundary. Performs EXIF timestamp freshness check (<5 min window) and anti-manipulation scoring.",
        "api": "POST /api/evidence/validate-location (verifies coordinate distance) and POST /api/evidence (stores evidence with liveness & anti-tamper scores)."
    },
    {
        "num": "7",
        "title": "CAQM GRAP (Graded Response Action Plan) Compliance Engine",
        "components": "shared/grapRules.ts, server/storage.ts, client/src/components/AuthorityDashboard.tsx",
        "concept": "Statutory rule-based compliance engine mapping live AQI to official Commission for Air Quality Management (CAQM) GRAP stages (Stage I: AQI 201-300, Stage II: 301-400, Stage III: 401-450, Stage IV: >450, No GRAP Required: <=200). Formulates ward-specific 7-day operational enforcement schedules anchored to plan generation dates (plan_generated_at).",
        "api": "shared/grapRules.ts (getOfficialGrapStage) served inside ward.intelligence_data via GET /api/wards."
    },
    {
        "num": "8",
        "title": "Interactive Policy Impact Simulator",
        "components": "server/policySimulator.ts, client/src/components/AuthorityDashboard.tsx",
        "concept": "Mathematical environmental impact projection model (simulatePolicy). Calculates immediate AQI reductions when municipal authorities adjust policy controls (Traffic Rationing %, Water Sprinkling %, Construction Halts), weighted by the ward's dominant pollution source.",
        "api": "POST /api/wards/:id/simulate-policy accepts { dustReduction, trafficReduction, constructionControl } and returns projected AQI, delta reduction, and mitigation breakdown."
    },
    {
        "num": "9",
        "title": "Multi-Source Live Telemetry Pipeline (WAQI + Open-Meteo)",
        "components": "server/storage.ts",
        "concept": "Dual-tier external air quality ingest. Fetches station telemetry from WAQI API across 290 mapped Delhi station UIDs and falls back to Open-Meteo Air Quality Grid API across 5 Delhi geographic zones (North, South, East, West, Central). Converts PM2.5 concentrations to AQI using CPCB NAQI standard piecewise linear equations.",
        "api": "WAQI API (api.waqi.info/feed/@${uid}/?token=${token}) and Open-Meteo Air Quality Grid API (air-quality-api.open-meteo.com/v1/air-quality)."
    },
    {
        "num": "10",
        "title": "Generative AI Ward Health Bulletin & Policy Advisor",
        "components": "client/src/components/NewsBulletin.tsx, server/routes.ts",
        "concept": "Generative AI natural language synthesis generating public health bulletins in English and Devanagari Hindi tailored to live ward statistics (AQI, PM2.5, PM10, dominant_source). Implements a 15-minute in-memory server cache (bulletinCache) to eliminate redundant LLM API calls.",
        "api": "Google Generative AI (gemini-1.5-flash) via GET /api/ward-bulletin (query params wardId, language) and GET /api/wards/:id/intelligence."
    },
    {
        "num": "11",
        "title": "Cigarette Health Risk Exposure Calculator",
        "components": "client/src/components/CigaretteHealthRiskCard.tsx",
        "concept": "Comparative health risk communication metric converting ambient PM2.5 concentrations into daily passive cigarette intake equivalents (~22 µg/m³ PM2.5 ≈ 1 cigarette/day). Displays 30-day cumulative cigarette exposure and comparative risk metrics.",
        "api": "Consumes pm25 and aqi telemetry from GET /api/wards."
    },
    {
        "num": "12",
        "title": "Bilingual Internationalization (i18n)",
        "components": "client/src/lib/i18n.ts",
        "concept": "React Context-based translation engine supporting instant runtime UI toggling between English and Devanagari Hindi (en / hi) across all dashboard widgets, navigation elements, forms, and AI prompts.",
        "api": "Pure client-side React Context (useLanguage hook)."
    }
]

for f in features:
    # Feature Heading
    h = doc.add_paragraph()
    h_run = h.add_run(f"Feature {f['num']}: {f['title']}")
    h_run.font.name = "Arial"
    h_run.font.size = Pt(14)
    h_run.font.bold = True
    h_run.font.color.rgb = RGBColor(16, 149, 193)
    
    # Table for structured details
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    labels = ["Component Files", "Concept & Mechanics", "APIs & Services Used"]
    vals = [f["components"], f["concept"], f["api"]]
    
    for i in range(3):
        row = table.rows[i]
        
        # Cell 0: Label
        cell_0 = row.cells[0]
        cell_0.width = Inches(1.8)
        p0 = cell_0.paragraphs[0]
        r0 = p0.add_run(labels[i])
        r0.font.name = "Arial"
        r0.font.size = Pt(10)
        r0.font.bold = True
        r0.font.color.rgb = RGBColor(30, 41, 59)
        
        # Cell 1: Value
        cell_1 = row.cells[1]
        cell_1.width = Inches(5.0)
        p1 = cell_1.paragraphs[0]
        r1 = p1.add_run(vals[i])
        r1.font.name = "Arial"
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(51, 65, 85)

    doc.add_paragraph() # Spacer between features

# Save in both workspace and artifacts folder
output_path_local = os.path.join(os.getcwd(), "Nirvayu_Project_Features_Documentation.docx")
output_path_artifact = r"C:\Users\DELL\.gemini\antigravity\brain\30d73cb3-ce38-4984-b566-dc57ab245022\Nirvayu_Project_Features_Documentation.docx"

doc.save(output_path_local)
doc.save(output_path_artifact)

print(f"Successfully saved Word Document to:\n1. {output_path_local}\n2. {output_path_artifact}")
